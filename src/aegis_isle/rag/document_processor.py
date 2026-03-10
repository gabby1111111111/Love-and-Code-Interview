"""
Document processing components for RAG pipeline.
Enhanced with multi-modal parsing, hybrid extraction strategies and OCR fallback.
"""

import hashlib
import mimetypes
import uuid
import time
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Union, Tuple

from pydantic import BaseModel, Field

from ..core.logging import logger


class ParsedTable(BaseModel):
    """Represents a parsed table from a document."""

    content: str  # Markdown formatted table
    position: int  # Position in document
    caption: Optional[str] = None
    metadata: Dict[str, Any] = Field(default_factory=dict)


class ParsedImage(BaseModel):
    """Represents a parsed image from a document."""

    description: Optional[str] = None
    position: int  # Position in document
    size: Optional[Tuple[int, int]] = None  # (width, height)
    format: Optional[str] = None  # png, jpg, etc.
    metadata: Dict[str, Any] = Field(default_factory=dict)


class EnhancedDocumentResult(BaseModel):
    """Enhanced document processing result with multi-modal support."""

    text: str
    images: List[ParsedImage] = Field(default_factory=list)
    tables: List[ParsedTable] = Field(default_factory=list)
    metadata: Dict[str, Any] = Field(default_factory=dict)


class DocumentMetadata(BaseModel):
    """Metadata for a document."""

    filename: str
    file_path: Optional[str] = None
    file_size: int = 0
    mime_type: Optional[str] = None
    language: Optional[str] = "en"
    source: str = "upload"  # upload, url, database, etc.
    author: Optional[str] = None
    title: Optional[str] = None
    created_at: datetime = Field(default_factory=datetime.now)
    modified_at: Optional[datetime] = None
    tags: List[str] = Field(default_factory=list)
    custom_fields: Dict[str, Any] = Field(default_factory=dict)
    extraction_method: Optional[str] = None  # fast_text, ocr, hybrid
    ocr_used: bool = False


class DocumentChunk(BaseModel):
    """Represents a chunk of a document."""

    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    document_id: str
    content: str
    chunk_index: int
    start_pos: int = 0
    end_pos: int = 0
    metadata: Dict[str, Any] = Field(default_factory=dict)
    embedding: Optional[List[float]] = None
    created_at: datetime = Field(default_factory=datetime.now)
    # New fields for enhanced chunking
    chunk_type: str = "text"  # text, table, image_description
    source_element: Optional[str] = None  # For tracking original table/image


class ProcessedDocument(BaseModel):
    """Represents a processed document."""

    id: str = Field(default_factory=lambda: str(uuid.uuid4()))
    content: str
    metadata: DocumentMetadata
    chunks: List[DocumentChunk] = Field(default_factory=list)
    processing_stats: Dict[str, Any] = Field(default_factory=dict)
    content_hash: Optional[str] = None
    created_at: datetime = Field(default_factory=datetime.now)

    def __post_init__(self):
        """Calculate content hash after initialization."""
        if not self.content_hash and self.content:
            self.content_hash = hashlib.sha256(
                self.content.encode('utf-8')
            ).hexdigest()


class EnhancedDocumentProcessor:
    """Enhanced document processor with hybrid parsing strategies and multi-modal support."""

    def __init__(
        self,
        enable_ocr: bool = True,
        text_threshold: int = 100,  # Minimum text length before switching to OCR
        enable_table_extraction: bool = True,
        enable_image_description: bool = True
    ):
        """Initialize enhanced document processor.

        Args:
            enable_ocr: Whether to enable OCR fallback
            text_threshold: Minimum text length before OCR fallback
            enable_table_extraction: Whether to extract tables
            enable_image_description: Whether to generate image descriptions
        """
        self.enable_ocr = enable_ocr
        self.text_threshold = text_threshold
        self.enable_table_extraction = enable_table_extraction
        self.enable_image_description = enable_image_description

        # Strategy mapping with priority (fast text first, OCR fallback)
        self.pdf_strategies = [
            self._process_pdf_fast,
            self._process_pdf_with_pdfplumber,
            self._process_pdf_with_ocr
        ]

        self.supported_types = {
            '.pdf': self._process_pdf_hybrid,
            '.docx': self._process_docx,
            '.doc': self._process_doc,
            '.txt': self._process_text,
            '.md': self._process_markdown,
            '.html': self._process_html,
            '.htm': self._process_html,
        }

        logger.info(
            f"Initialized Enhanced Document Processor - "
            f"OCR: {enable_ocr}, Table extraction: {enable_table_extraction}"
        )

    async def process_file(
        self,
        file_path: Union[str, Path],
        metadata: Optional[Dict[str, Any]] = None
    ) -> ProcessedDocument:
        """Process a file with enhanced multi-modal extraction."""
        file_path = Path(file_path)

        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        logger.info(f"Processing file with enhanced processor: {file_path}")

        # Determine file type
        file_ext = file_path.suffix.lower()
        if file_ext not in self.supported_types:
            raise ValueError(f"Unsupported file type: {file_ext}")

        # Create metadata
        doc_metadata = self._create_metadata(file_path, metadata or {})

        # Process content with enhanced strategy
        start_time = time.time()
        processor = self.supported_types[file_ext]
        enhanced_result = await processor(file_path)

        processing_time = time.time() - start_time
        enhanced_result.metadata['processing_time'] = processing_time

        # Update document metadata
        doc_metadata.extraction_method = enhanced_result.metadata.get('extraction_method', 'unknown')
        doc_metadata.ocr_used = enhanced_result.metadata.get('ocr_used', False)

        # Create processed document with enhanced content
        content = enhanced_result.text

        # Append table information to content
        if enhanced_result.tables:
            table_content = "\n\n".join([
                f"Table {i+1}:\n{table.content}"
                for i, table in enumerate(enhanced_result.tables)
            ])
            content += f"\n\n=== TABLES ===\n{table_content}"

        # Append image descriptions to content
        if enhanced_result.images:
            image_content = "\n\n".join([
                f"Image {i+1}: {img.description or 'No description available'}"
                for i, img in enumerate(enhanced_result.images) if img.description
            ])
            if image_content:
                content += f"\n\n=== IMAGES ===\n{image_content}"

        document = ProcessedDocument(
            content=content,
            metadata=doc_metadata,
            processing_stats=enhanced_result.metadata
        )

        # Store enhanced data in processing stats for chunker access
        document.processing_stats.update({
            'tables': [table.dict() for table in enhanced_result.tables],
            'images': [img.dict() for img in enhanced_result.images],
            'enhanced_result': enhanced_result.dict()
        })

        logger.info(
            f"Enhanced processing completed: {len(content)} chars, "
            f"{len(enhanced_result.tables)} tables, {len(enhanced_result.images)} images, "
            f"{processing_time:.2f}s"
        )

        return document

    async def _process_pdf_hybrid(self, file_path: Path) -> EnhancedDocumentResult:
        """Process PDF with hybrid strategy: fast text first, then OCR fallback."""
        logger.info(f"Processing PDF with hybrid strategy: {file_path}")

        for i, strategy in enumerate(self.pdf_strategies):
            try:
                result = await strategy(file_path)

                # Check if we got enough text
                if len(result.text.strip()) >= self.text_threshold:
                    strategy_name = ['fast_text', 'pdfplumber', 'ocr'][i]
                    result.metadata['extraction_method'] = strategy_name
                    result.metadata['strategy_index'] = i
                    logger.info(f"PDF processing succeeded with {strategy_name} strategy")
                    return result
                else:
                    logger.debug(f"Strategy {i} produced insufficient text ({len(result.text)} chars)")

            except Exception as e:
                logger.warning(f"PDF processing strategy {i} failed: {e}")
                continue

        # All strategies failed
        logger.error(f"All PDF processing strategies failed for {file_path}")
        return EnhancedDocumentResult(
            text="Failed to extract text from PDF",
            metadata={'extraction_method': 'failed', 'error': 'All strategies failed'}
        )

    async def _process_pdf_fast(self, file_path: Path) -> EnhancedDocumentResult:
        """Fast PDF text extraction using PyPDF2."""
        start_time = time.time()

        try:
            import PyPDF2

            with open(file_path, 'rb') as file:
                reader = PyPDF2.PdfReader(file)
                content_parts = []

                for page_num, page in enumerate(reader.pages):
                    text = page.extract_text()
                    if text.strip():
                        content_parts.append(text)

                content = "\n\n".join(content_parts)

                return EnhancedDocumentResult(
                    text=content,
                    metadata={
                        "total_pages": len(reader.pages),
                        "processing_time": time.time() - start_time,
                        "extraction_method": "fast_text",
                        "library": "PyPDF2"
                    }
                )

        except Exception as e:
            logger.debug(f"Fast PDF extraction failed: {e}")
            raise

    async def _process_pdf_with_pdfplumber(self, file_path: Path) -> EnhancedDocumentResult:
        """Enhanced PDF processing with pdfplumber for better text and table extraction."""
        start_time = time.time()

        try:
            import pdfplumber

            text_parts = []
            tables = []
            images = []

            with pdfplumber.open(str(file_path)) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    # Extract text
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)

                    # Extract tables if enabled
                    if self.enable_table_extraction:
                        page_tables = page.extract_tables()
                        for table_idx, table_data in enumerate(page_tables):
                            if table_data:
                                # Convert table to markdown
                                markdown_table = self._table_to_markdown(table_data)
                                table = ParsedTable(
                                    content=markdown_table,
                                    position=page_num,
                                    metadata={
                                        'page': page_num,
                                        'table_index': table_idx,
                                        'rows': len(table_data),
                                        'cols': len(table_data[0]) if table_data else 0
                                    }
                                )
                                tables.append(table)

                    # Extract images (basic info only)
                    if self.enable_image_description and hasattr(page, 'images'):
                        for img_idx, img in enumerate(page.images):
                            parsed_img = ParsedImage(
                                description=f"Image on page {page_num + 1}",
                                position=page_num,
                                metadata={
                                    'page': page_num,
                                    'image_index': img_idx
                                }
                            )
                            images.append(parsed_img)

            content = "\n\n".join(text_parts)

            return EnhancedDocumentResult(
                text=content,
                tables=tables,
                images=images,
                metadata={
                    "total_pages": len(pdf.pages),
                    "processing_time": time.time() - start_time,
                    "extraction_method": "pdfplumber",
                    "tables_extracted": len(tables),
                    "images_found": len(images)
                }
            )

        except Exception as e:
            logger.debug(f"pdfplumber extraction failed: {e}")
            raise

    async def _process_pdf_with_ocr(self, file_path: Path) -> EnhancedDocumentResult:
        """Process PDF with OCR for scanned documents."""
        if not self.enable_ocr:
            raise ValueError("OCR is disabled")

        start_time = time.time()

        try:
            import pdf2image
            import pytesseract
            from PIL import Image

            # Convert PDF to images
            images_list = pdf2image.convert_from_path(str(file_path))
            content_parts = []
            parsed_images = []

            for page_num, image in enumerate(images_list):
                # Extract text using OCR
                text = pytesseract.image_to_string(image, lang='eng+chi_sim')
                if text.strip():
                    content_parts.append(text)

                # Create image description if enabled
                if self.enable_image_description:
                    img_desc = f"Scanned page {page_num + 1} from PDF"
                    parsed_img = ParsedImage(
                        description=img_desc,
                        position=page_num,
                        size=(image.width, image.height),
                        format='PIL_Image',
                        metadata={
                            'page': page_num,
                            'ocr_extracted': True,
                            'text_length': len(text)
                        }
                    )
                    parsed_images.append(parsed_img)

            content = "\n\n".join(content_parts)

            return EnhancedDocumentResult(
                text=content,
                images=parsed_images,
                metadata={
                    "total_pages": len(images_list),
                    "processing_time": time.time() - start_time,
                    "extraction_method": "ocr",
                    "ocr_used": True,
                    "ocr_language": "eng+chi_sim",
                    "images_processed": len(images_list)
                }
            )

        except Exception as e:
            logger.error(f"OCR processing failed for {file_path}: {e}")
            raise

    def _table_to_markdown(self, table_data: List[List[str]]) -> str:
        """Convert table data to markdown format."""
        if not table_data or not table_data[0]:
            return ""

        markdown_rows = []

        # Header row
        header = table_data[0]
        markdown_rows.append("| " + " | ".join(str(cell or "") for cell in header) + " |")

        # Separator row
        separator = "|" + "|".join(" --- " for _ in header) + "|"
        markdown_rows.append(separator)

        # Data rows
        for row in table_data[1:]:
            if row:  # Skip empty rows
                markdown_row = "| " + " | ".join(str(cell or "") for cell in row) + " |"
                markdown_rows.append(markdown_row)

        return "\n".join(markdown_rows)

    def _create_metadata(
        self,
        file_path: Path,
        custom_metadata: Dict[str, Any]
    ) -> DocumentMetadata:
        """Create enhanced metadata for a file."""
        stat = file_path.stat()
        mime_type, _ = mimetypes.guess_type(str(file_path))
        custom_metadata.pop('filename', None)
        custom_metadata.pop('file_path', None)
        custom_metadata.pop('file_size', None)
        return DocumentMetadata(
            filename=file_path.name,
            file_path=str(file_path),
            file_size=stat.st_size,
            mime_type=mime_type,
            created_at=datetime.fromtimestamp(stat.st_ctime),
            modified_at=datetime.fromtimestamp(stat.st_mtime),
            **custom_metadata
        )


# Legacy compatibility class
class DocumentProcessor(EnhancedDocumentProcessor):
    """Main document processor that handles different file types.

    Legacy wrapper around EnhancedDocumentProcessor for backward compatibility.
    """

    def __init__(self):
        """Initialize with default settings for backward compatibility."""
        super().__init__(
            enable_ocr=True,
            text_threshold=100,
            enable_table_extraction=True,
            enable_image_description=True
        )

        # Legacy attribute for compatibility
        self.ocr_enabled = self.enable_ocr

    async def process_text(
        self,
        content: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> ProcessedDocument:
        """Process raw text content."""
        doc_metadata = DocumentMetadata(
            filename="text_input",
            mime_type="text/plain",
            source="text_input",
            extraction_method="direct_text",
            **(metadata or {})
        )

        document = ProcessedDocument(
            content=content,
            metadata=doc_metadata,
            processing_stats={"content_length": len(content)}
        )

        return document

    async def process_url(
        self,
        url: str,
        metadata: Optional[Dict[str, Any]] = None
    ) -> ProcessedDocument:
        """Process content from a URL."""
        import aiohttp

        try:
            async with aiohttp.ClientSession() as session:
                async with session.get(url) as response:
                    content = await response.text()

            doc_metadata = DocumentMetadata(
                filename=url.split('/')[-1] or "webpage",
                source="url",
                mime_type=response.headers.get('content-type', 'text/html'),
                extraction_method="web_scraping",
                **(metadata or {}),
                custom_fields={"url": url}
            )

            # Process HTML content
            if 'html' in doc_metadata.mime_type:
                content, stats = await self._process_html_content(content)
            else:
                stats = {"content_length": len(content)}

            document = ProcessedDocument(
                content=content,
                metadata=doc_metadata,
                processing_stats=stats
            )

            return document

        except Exception as e:
            logger.error(f"Error processing URL {url}: {e}")
            raise

    # Legacy methods that need to be implemented for compatibility
    async def _process_docx(self, file_path: Path) -> EnhancedDocumentResult:
        """Process DOCX file."""
        start_time = time.time()

        try:
            from docx import Document

            doc = Document(str(file_path))
            content_parts = []
            tables = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content_parts.append(paragraph.text)

            # Process tables with enhanced extraction
            for table_idx, table in enumerate(doc.tables):
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)

                if table_data:
                    markdown_table = self._table_to_markdown(table_data)
                    parsed_table = ParsedTable(
                        content=markdown_table,
                        position=table_idx,
                        metadata={
                            'table_index': table_idx,
                            'rows': len(table_data),
                            'cols': len(table_data[0]) if table_data else 0
                        }
                    )
                    tables.append(parsed_table)

            content = "\n".join(content_parts)

            return EnhancedDocumentResult(
                text=content,
                tables=tables,
                metadata={
                    "total_paragraphs": len(doc.paragraphs),
                    "total_tables": len(tables),
                    "processing_time": time.time() - start_time,
                    "extraction_method": "docx"
                }
            )

        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {e}")
            raise

    async def _process_doc(self, file_path: Path) -> EnhancedDocumentResult:
        """Process DOC file (legacy Word format)."""
        # For .doc files, we'd typically use python-docx2txt or similar
        raise NotImplementedError(
            "Legacy .doc format not supported. Please convert to .docx"
        )

    async def _process_text(self, file_path: Path) -> EnhancedDocumentResult:
        """Process plain text file."""
        start_time = time.time()

        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read()

            return EnhancedDocumentResult(
                text=content,
                metadata={
                    "content_length": len(content),
                    "processing_time": time.time() - start_time,
                    "extraction_method": "text",
                    "encoding": "utf-8"
                }
            )

        except UnicodeDecodeError:
            # Try different encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        content = file.read()

                    return EnhancedDocumentResult(
                        text=content,
                        metadata={
                            "content_length": len(content),
                            "processing_time": time.time() - start_time,
                            "extraction_method": "text",
                            "encoding": encoding
                        }
                    )
                except UnicodeDecodeError:
                    continue

            raise ValueError(f"Unable to decode text file: {file_path}")

    async def _process_markdown(self, file_path: Path) -> EnhancedDocumentResult:
        """Process Markdown file."""
        return await self._process_text(file_path)

    async def _process_html(self, file_path: Path) -> EnhancedDocumentResult:
        """Process HTML file."""
        start_time = time.time()

        with open(file_path, 'r', encoding='utf-8') as file:
            html_content = file.read()

        content, additional_stats = await self._process_html_content(html_content)

        return EnhancedDocumentResult(
            text=content,
            metadata={
                "processing_time": time.time() - start_time,
                "extraction_method": "html",
                **additional_stats
            }
        )

    async def _process_html_content(self, html_content: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from HTML content."""
        try:
            from bs4 import BeautifulSoup

            soup = BeautifulSoup(html_content, 'html.parser')

            # Remove script and style elements
            for script in soup(["script", "style"]):
                script.decompose()

            # Get text content
            text = soup.get_text()

            # Clean up text
            lines = (line.strip() for line in text.splitlines())
            chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
            content = "\n".join(chunk for chunk in chunks if chunk)

            stats = {
                "original_length": len(html_content),
                "extracted_length": len(content),
                "extraction_method": "BeautifulSoup"
            }

            return content, stats

        except ImportError:
            logger.warning("BeautifulSoup not available, using regex fallback")
            import re

            # Simple regex-based HTML tag removal
            clean = re.compile('<.*?>')
            content = re.sub(clean, '', html_content)

            stats = {
                "original_length": len(html_content),
                "extracted_length": len(content),
                "extraction_method": "regex"
            }

            return content, stats